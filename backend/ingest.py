"""Ingestion: load PDF/DOCX, semantic chunking, parent-child hierarchy."""
import base64
import hashlib
import json
import re
import unicodedata
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Any, Callable, Optional

import tiktoken
from docling.document_converter import DocumentConverter 
from langchain_core.documents import Document
from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_voyageai import VoyageAIEmbeddings
from config import get_settings, get_runtime_embedding_settings
from legal_tokenizer import legal_tokenize as _legal_tokenize
from pg_client import (
    insert_sparse_vocab_entries,
    load_sparse_vocab_map,
    upsert_sparse_bm25_stats,
    upsert_vector_parents,
    insert_vector_chunks,
    get_collection_count,
)
from utils import fix_position_ids as _fix_position_ids, tokenize_for_sparse as _tokenize_for_sparse


TABLE_START = "<!--TABLE_START-->"
TABLE_END = "<!--TABLE_END-->"



def _sanitize_metadata_value(value: Any) -> Any:
    if isinstance(value, (str, int, float, bool)) or value is None:
        return value
    if isinstance(value, (list, tuple)):
        cleaned: list[Any] = []
        for item in value:
            cleaned.append(_sanitize_metadata_value(item))
        return cleaned
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _sanitize_text_for_api(text: str) -> str:
    """Normalize text to be safe for JSON APIs (remove NUL and invalid UTF-8 surrogates)."""
    if not isinstance(text, str):
        text = str(text)
    cleaned = text.replace("\x00", "")
    cleaned = cleaned.encode("utf-8", errors="ignore").decode("utf-8", errors="ignore")
    return "".join(
        ch for ch in cleaned
        if (ch in "\n\r\t") or (unicodedata.category(ch) not in {"Cc", "Cs"})
    )



# --- Sparse vector helpers (for Qdrant native sparse) ---


def _build_vocab(
    texts: list[str],
    max_vocab_size: int = 100_000,
) -> tuple[dict[str, int], dict[str, float], float]:
    """Build vocabulary and BM25 statistics from corpus.

    Returns:
        vocab:    token → index mapping (top-N by frequency)
        idf_map:  token → Robertson IDF value
        avgdl:    average document length in tokens
    """
    from collections import Counter
    from math import log

    doc_token_lists: list[list[str]] = [_legal_tokenize(t) for t in texts]
    N = len(doc_token_lists)

    corpus_counter: Counter[str] = Counter()
    df_counter: Counter[str] = Counter()
    total_tokens = 0
    for token_list in doc_token_lists:
        corpus_counter.update(token_list)
        df_counter.update(set(token_list))
        total_tokens += len(token_list)

    avgdl = total_tokens / N if N > 0 else 1.0

    vocab: dict[str, int] = {}
    for token, _ in corpus_counter.most_common(max_vocab_size):
        vocab[token] = len(vocab)

    # Robertson IDF: log((N - df + 0.5) / (df + 0.5) + 1)
    idf_map: dict[str, float] = {}
    for token in vocab:
        df = df_counter.get(token, 0)
        idf_map[token] = log((N - df + 0.5) / (df + 0.5) + 1.0)

    return vocab, idf_map, avgdl


def _text_to_sparse_vector(
    text: str,
    vocab: dict[str, int],
    idf_map: dict[str, float] | None = None,
    avgdl: float = 1.0,
    k1: float = 1.5,
    b: float = 0.75,
) -> tuple[list[int], list[float]]:
    """Convert text to (indices, values) for Qdrant SparseVector using BM25.

    When ``idf_map`` is provided, applies the full Okapi BM25 formula with
    document-length normalization.  Falls back to ``1 + log(tf)`` when
    ``idf_map`` is None or empty.
    """
    from math import log
    tokens = _legal_tokenize(text)
    if not tokens:
        return [], []
    dl = len(tokens)
    tf_raw: dict[int, int] = {}
    for t in tokens:
        idx = vocab.get(t)
        if idx is None:
            continue
        tf_raw[idx] = tf_raw.get(idx, 0) + 1

    scores: dict[int, float] = {}
    if idf_map:
        idx_to_token = {v: k for k, v in vocab.items()}
        for idx, raw_tf in tf_raw.items():
            token = idx_to_token.get(idx, "")
            idf = idf_map.get(token, 1.0)
            tf_norm = (raw_tf * (k1 + 1)) / (raw_tf + k1 * (1 - b + b * dl / avgdl))
            scores[idx] = idf * tf_norm
    else:
        for idx, raw_tf in tf_raw.items():
            scores[idx] = 1.0 + log(raw_tf)

    indices = sorted(scores.keys())
    values = [float(scores[i]) for i in indices]
    return indices, values


# --- Tokenization & text splitting ---

def get_encoding():
    return tiktoken.get_encoding("cl100k_base")


def count_tokens(text: str) -> int:
    return len(get_encoding().encode(text))


def chunk_by_tokens(text: str, max_tokens: int, overlap: int = 50) -> list[str]:
    """Split text into chunks by token count with overlap."""
    enc = get_encoding()
    tokens = enc.encode(text)
    out = []
    start = 0
    while start < len(tokens):
        end = min(start + max_tokens, len(tokens))
        out.append(enc.decode(tokens[start:end]))
        start = end - overlap if end < len(tokens) else len(tokens)
    return out


def _split_paragraphs(text: str) -> list[str]:
    """Split text into paragraphs, keeping TABLE_START...TABLE_END blocks atomic."""
    table_pattern = re.compile(
        rf"({re.escape(TABLE_START)}.*?{re.escape(TABLE_END)})",
        re.DOTALL,
    )
    segments = table_pattern.split(text)
    parts: list[str] = []
    for seg in segments:
        seg_stripped = seg.strip()
        if not seg_stripped:
            continue
        if seg_stripped.startswith(TABLE_START):
            parts.append(seg_stripped)
        else:
            for p in re.split(r"\n\s*\n+", seg):
                p = p.strip()
                if p:
                    parts.append(p)
    if not parts and text.strip():
        parts = [text.strip()]
    return parts


def _split_sentences(text: str) -> list[str]:
    """Simple rule-based sentence splitter (Vietnamese and English)."""
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return []
    parts = re.split(r"(?<=[.!?])\s+", cleaned)
    return [p.strip() for p in parts if p.strip()]


def _split_words(text: str) -> list[str]:
    return text.split()


# --- Legal article parsing ---

# Matches Vietnamese and English article markers at the start of a line.
# Group 1: keyword (Điều/Article/Section), Group 2: number, Group 3: optional title text.
_ARTICLE_RE = re.compile(
    r"(?im)^(điều|article|section)\s+(\d+)[.:]?\s*(.*)",
)

# Matches chapter/part markers.
# Group 1: keyword, Group 2: Roman or Arabic numeral, Group 3: optional title text.
_CHAPTER_RE = re.compile(
    r"(?im)^(chương|chapter|part|phần)\s+([IVXLCDM]+|\d+)[.:]?\s*(.*)",
)

# Matches numbered clause lines inside an article, e.g. "1. " or "  2. "
_CLAUSE_RE = re.compile(r"(?m)^[ \t]*(\d+)\.\s+")


def _detect_legal_articles(text: str) -> bool:
    """Return True if *text* contains at least 5 article markers (Điều/Article/Section)."""
    return len(_ARTICLE_RE.findall(text)) >= 5


def _split_legal_articles(text: str) -> list[dict]:
    """Parse *text* into a list of article dicts.

    Each dict has:
    - ``article_number`` (str): e.g. "15"
    - ``article_title``  (str): heading text after "Điều 15."
    - ``chapter_number`` (str): current chapter numeral, e.g. "II"
    - ``chapter_title``  (str): current chapter heading text
    - ``body``           (str): full article text including its header line

    Text before the first article is returned as a preamble entry with
    ``article_number="0"``.
    """
    lines = text.splitlines(keepends=True)

    current_chapter_num = ""
    current_chapter_title = ""

    articles: list[dict] = []
    buf: list[str] = []
    current_article_num = "0"
    current_article_title = ""
    current_chapter_at_start = ""
    current_chapter_title_at_start = ""

    def _flush(buf, art_num, art_title, chap_num, chap_title):
        body = "".join(buf).strip()
        if body:
            articles.append({
                "article_number": art_num,
                "article_title": art_title,
                "chapter_number": chap_num,
                "chapter_title": chap_title,
                "body": body,
            })

    for line in lines:
        stripped = line.strip()

        chap_m = _CHAPTER_RE.match(stripped)
        if chap_m:
            _flush(buf, current_article_num, current_article_title,
                   current_chapter_at_start, current_chapter_title_at_start)
            buf = [line]
            current_chapter_num = chap_m.group(2).strip()
            current_chapter_title = chap_m.group(3).strip()
            current_article_num = "0"
            current_article_title = ""
            current_chapter_at_start = current_chapter_num
            current_chapter_title_at_start = current_chapter_title
            continue

        art_m = _ARTICLE_RE.match(stripped)
        if art_m:
            _flush(buf, current_article_num, current_article_title,
                   current_chapter_at_start, current_chapter_title_at_start)
            buf = [line]
            current_article_num = art_m.group(2).strip()
            current_article_title = art_m.group(3).strip()
            current_chapter_at_start = current_chapter_num
            current_chapter_title_at_start = current_chapter_title
            continue

        buf.append(line)

    _flush(buf, current_article_num, current_article_title,
           current_chapter_at_start, current_chapter_title_at_start)

    return articles


def _split_clauses(text: str) -> list[str]:
    """Split article text on numbered clause markers (e.g. '1. ', '2. ').

    Falls back to sentence splitting when fewer than 2 clauses are found.
    """
    matches = list(_CLAUSE_RE.finditer(text))
    if len(matches) < 2:
        return _split_sentences(text)

    chunks: list[str] = []
    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
    return chunks


def _load_pdf_docling(path: Path, settings=None) -> list[Document]:
    """
    Load PDF via Docling and convert it to rich Markdown text for RAG.

    - Uses Docling's DocumentConverter to parse the PDF.
    - Exports the full document to Markdown (headings, lists, tables).
    - Ignores images (no GPT-4 Vision).
    """
    if settings is None:
        settings = get_settings()

    converter = DocumentConverter()
    try:
        # Docling auto-detects the format from the file path
        result = converter.convert(str(path))
    except Exception:
        # On failure, return empty and let caller handle the error
        return []

    doc = getattr(result, "document", None)
    if doc is None:
        return []

    try:
        markdown = doc.export_to_markdown()
    except Exception:
        return []

    markdown = (markdown or "").strip()
    if not markdown:
        return []

    return [
        Document(
            page_content=markdown,
            metadata={
                "source": path.name,
                "file_path": str(path),
                # Treat entire Docling output as a single logical page/section
                "page": 0,
            },
        )
    ]


def _load_docx(path: Path) -> list[Document]:
    """Load DOCX via python-docx (no extra system deps)."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        try:
            from langchain_community.document_loaders import UnstructuredWordDocumentLoader
            loader = UnstructuredWordDocumentLoader(str(path))
            return loader.load()
        except Exception:
            raise ImportError("Install python-docx: pip install python-docx")
    doc = DocxDocument(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n\n".join(parts) or ""
    if not text.strip():
        return []
    return [Document(page_content=text, metadata={"source": path.name, "file_path": str(path)})]


def load_document(file_path: Path, settings=None) -> list[Document]:
    """Load a single PDF or DOCX file into LangChain documents."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        docs = _load_pdf_docling(path, settings=settings)
    elif suffix in (".docx", ".doc"):
        docs = _load_docx(path)
    else:
        raise ValueError(f"Unsupported format: {suffix}. Use .pdf or .docx")

    for d in docs:
        d.metadata.setdefault("source", str(path.name))
        d.metadata["file_path"] = str(path)
    return docs


# --- Parent and child chunking ---

_PRESPLIT_THRESHOLD_TOKENS = 50_000
_PRESPLIT_SECTION_TOKENS = 10_000


def _presplit_large_doc(doc: Document) -> list[Document]:
    """Split a very large document into sections of ~_PRESPLIT_SECTION_TOKENS
    on paragraph boundaries so SemanticChunker doesn't choke."""
    total = count_tokens(doc.page_content)
    if total <= _PRESPLIT_THRESHOLD_TOKENS:
        return [doc]

    paras = _split_paragraphs(doc.page_content)
    sections: list[Document] = []
    buf: list[str] = []
    buf_tokens = 0
    for p in paras:
        n = count_tokens(p)
        if buf_tokens + n > _PRESPLIT_SECTION_TOKENS and buf:
            sections.append(
                Document(page_content="\n\n".join(buf), metadata=dict(doc.metadata))
            )
            buf = [p]
            buf_tokens = n
        else:
            buf.append(p)
            buf_tokens += n
    if buf:
        sections.append(
            Document(page_content="\n\n".join(buf), metadata=dict(doc.metadata))
        )
    return sections


def semantic_parent_chunks(
    documents: list[Document],
    parent_target_tokens: int,
    embeddings: Any,
) -> list[Document]:
    """
    Split markdown documents by their headers, then recursively character split
    any remaining overly long sections to fit the target token limit.
    This replaces the expensive SemanticChunker.
    """
    headers_to_split_on = [
        ("#", "Header 1"),
        ("##", "Header 2"),
        ("###", "Header 3"),
    ]
    markdown_splitter = MarkdownHeaderTextSplitter(
        headers_to_split_on=headers_to_split_on,
        strip_headers=False,
    )
    
    char_splitter = RecursiveCharacterTextSplitter(
        chunk_size=parent_target_tokens * 4,  # Approx 4 chars per token
        chunk_overlap=200,
        separators=["\n\n", "\n", ".", " ", ""],
    )

    parents: list[Document] = []
    for doc in documents:
        text = doc.page_content or ""
        md_splits = markdown_splitter.split_text(text)
        
        for md_split in md_splits:
            # Reattach original file-level metadata to the markdown sections
            merged_metadata = {**doc.metadata, **md_split.metadata}
            md_split.metadata = merged_metadata
            
            # If the md section is still too big, recursive character split it
            if count_tokens(md_split.page_content) > parent_target_tokens:
                sub_splits = char_splitter.split_documents([md_split])
                parents.extend(sub_splits)
            else:
                parents.append(md_split)

    return parents


def article_parent_chunks(
    documents: list[Document],
    parent_target_tokens: int,
    embeddings: Any,
) -> list[Document]:
    """Split legal documents into parent chunks, one per article (Điều/Article/Section).

    Each article is kept as a single parent chunk.  If an article exceeds
    *parent_target_tokens*, it is further split with RecursiveCharacterTextSplitter
    while preserving the article-level metadata on every sub-chunk.

    Falls back to :func:`semantic_parent_chunks` when no article markers are found.
    """
    char_splitter = RecursiveCharacterTextSplitter(
        chunk_size=parent_target_tokens * 4,
        chunk_overlap=100,
        separators=["\n\n", "\n", ".", " ", ""],
    )

    parents: list[Document] = []
    for doc in documents:
        text = doc.page_content or ""
        if not _detect_legal_articles(text):
            parents.extend(
                semantic_parent_chunks([doc], parent_target_tokens, embeddings)
            )
            continue

        article_dicts = _split_legal_articles(text)
        for art in article_dicts:
            if art["article_number"] == "0" and len(art["body"].split()) < 30:
                continue
            article_meta = {
                **doc.metadata,
                "article_number": art["article_number"],
                "article_title": art["article_title"],
                "chapter_number": art["chapter_number"],
                "chapter_title": art["chapter_title"],
                "chunk_strategy": "article",
            }
            if count_tokens(art["body"]) > parent_target_tokens:
                article_doc = Document(page_content=art["body"], metadata=article_meta)
                sub_splits = char_splitter.split_documents([article_doc])
                parents.extend(sub_splits)
            else:
                parents.append(Document(page_content=art["body"], metadata=article_meta))

    return parents


def build_parent_chunks(
    documents: list[Document],
    settings,
    embeddings: Any,
) -> list[Document]:
    """Build parent chunks using article-based chunking.

    Each Điều/Article/Section becomes one parent chunk.  Documents without
    article markers fall back to semantic (header-based) splitting.
    """
    return article_parent_chunks(
        documents,
        parent_target_tokens=getattr(settings, "article_max_parent_tokens", 1200),
        embeddings=embeddings,
    )


def parent_to_children_dynamic(parent: Document, settings, *, effective_strategy: Optional[str] = None) -> list[Document]:
    """Split a parent chunk into child chunks by clause (Khoản).

    Splits on numbered clause markers (1. 2. …); falls back to sentences
    when fewer than 2 clauses are found.
    """
    text = parent.page_content or ""
    chunks = _split_clauses(text)

    article_no = parent.metadata.get("article_number", "")
    article_title = parent.metadata.get("article_title", "")

    children: list[Document] = []
    for i, c in enumerate(chunks):
        content = c.strip()
        if not content:
            continue

        # Prepend article + khoản context so each child chunk is self-contained.
        # Only applies to legal article chunks where a numbered clause ("1. " "2. ")
        # is detected — sentence-split fallback chunks are left unchanged.
        clause_match = _CLAUSE_RE.match(content)
        if clause_match and article_no:
            clause_no = clause_match.group(1)
            header = (
                f"Điều {article_no}, khoản {clause_no} ({article_title}): "
                if article_title
                else f"Điều {article_no}, khoản {clause_no}: "
            )
            content = header + content

        children.append(
            Document(
                page_content=content,
                metadata={
                    **parent.metadata,
                    "chunk_index": i,
                    "clause_number": clause_match.group(1) if clause_match else None,
                    "parent_content": text,
                },
            )
        )
    return children


# --- Collection name ---

def sanitize_collection_name(name: str) -> str:
    """Normalize collection name: 3-63 chars, alphanumeric/underscore/hyphen."""
    if not name or not str(name).strip():
        return "rag_chatbot"
    s = str(name).strip()
    s = re.sub(r"[^a-zA-Z0-9_-]", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    if not s:
        return "rag_chatbot"
    if len(s) < 3:
        s = s + "xx"[: 3 - len(s)]
    if not s[0].isalnum():
        s = "c" + s
    if not s[-1].isalnum():
        s = s + "1"
    return s[:63]


def _noop_progress(step: str, msg: str, current: int = 0, total: int = 0) -> None:
    pass


# --- Parallel embedding ---

def _parallel_embed(
    embeddings,
    texts: list[str],
    batch_size: int = 64,
    max_workers: int = 4,
    progress_fn: Optional[Callable[[str, str, int, int], None]] = None,
) -> list[list[float]]:
    """Embed texts in parallel batches using a thread pool."""
    if not texts:
        return []

    texts = [_sanitize_text_for_api(t) for t in texts]
    batches = [texts[i : i + batch_size] for i in range(0, len(texts), batch_size)]
    total_batches = len(batches)

    # OpenAI embeddings client is not safely parallelized in this flow and may emit malformed requests.
    # Force sequential batching for OpenAI models to avoid intermittent 400 invalid JSON errors.
    model_name = str(getattr(embeddings, "model", "") or "")
    if model_name.startswith("text-embedding"):
        max_workers = 1

    if total_batches <= 1:
        if progress_fn:
            progress_fn("embed", f"Embedding {len(texts)} chunks...", 0, 1)
        result = embeddings.embed_documents(texts)
        if progress_fn:
            progress_fn("embed", f"Embedded {len(result)} chunks", 1, 1)
        return result

    all_embeddings: list[Optional[list[list[float]]]] = [None] * total_batches
    completed = 0

    if progress_fn:
        progress_fn("embed", f"Embedding {len(texts)} chunks in {total_batches} batches...", 0, total_batches)

    def _embed_batch(idx: int, batch: list[str]) -> tuple[int, list[list[float]]]:
        import time as _time
        for attempt in range(5):
            try:
                return idx, embeddings.embed_documents(batch)
            except Exception as e:
                msg = str(e).lower()
                if "rate_limit" in msg or "429" in msg:
                    wait = 2 ** attempt
                    _time.sleep(wait)
                    continue
                if "parse the json body" in msg or "not valid json" in msg:
                    fixed_batch = [_sanitize_text_for_api(item) for item in batch]
                    _time.sleep(1)
                    try:
                        return idx, embeddings.embed_documents(fixed_batch)
                    except Exception:
                        raise
                raise
        return idx, embeddings.embed_documents(batch)

    with ThreadPoolExecutor(max_workers=min(max_workers, total_batches)) as pool:
        futures = {
            pool.submit(_embed_batch, i, batch): i
            for i, batch in enumerate(batches)
        }
        for future in as_completed(futures):
            idx, batch_result = future.result()
            all_embeddings[idx] = batch_result
            completed += 1
            if progress_fn:
                progress_fn("embed", f"Batch {completed}/{total_batches} done", completed, total_batches)

    result: list[list[float]] = []
    for batch_result in all_embeddings:
        result.extend(batch_result)
    return result


def _insert_chunks_in_batches(
    collection_name: str,
    records: list[dict],
    batch_size: int,
    progress_fn: Optional[Callable[[str, str, int, int], None]] = None,
) -> None:
    if not records:
        return
    effective_batch = max(1, int(batch_size or 128))
    total = len(records)
    batches = (total + effective_batch - 1) // effective_batch
    for i in range(batches):
        start = i * effective_batch
        end = min(start + effective_batch, total)
        if progress_fn:
            progress_fn("pgvector", f"Writing chunks {start + 1}-{end}/{total}...", i, batches)
        insert_vector_chunks(collection_name, records[start:end])
    if progress_fn:
        progress_fn("pgvector", f"Writing chunks {total}/{total} complete", batches, batches)


# --- Full file ingestion ---

def ingest_documents(
    documents: list[Document],
    *,
    source_label: str,
    collection_name: str = "rag_chatbot",
    on_progress: Optional[Callable[[str, str, int, int], None]] = None,
) -> dict[str, Any]:
    """Ingest pre-loaded documents into the parent-child pipeline."""
    progress = on_progress or _noop_progress
    settings = get_settings()
    if not settings.openai_api_key:
        raise ValueError("OPENAI_API_KEY is required. Set it in .env or environment.")

    if not documents:
        return {"error": "No content extracted", "file": source_label}

    progress("load", f"Loaded {len(documents)} page(s)", 1, 1)

    emb_rt = get_runtime_embedding_settings()
    model = emb_rt.get("model") or "voyage-law-2"
    if model.startswith("voyage-"):
        embeddings = VoyageAIEmbeddings(
            model=model,
            voyage_api_key=emb_rt.get("voyage_api_key") or None,
        )
    elif model.startswith("text-embedding"):
        embeddings = OpenAIEmbeddings(
            model=model,
            api_key=emb_rt.get("openai_api_key") or "",
        )
    else:
        embeddings = HuggingFaceEmbeddings(
            model_name=model,
            model_kwargs={"trust_remote_code": True, "device": "cpu"},
        )
        _fix_position_ids(embeddings.client)

    progress("semantic", "Chunking (article)...", 0, 0)
    effective_strategy = "article"

    parents = build_parent_chunks(
        documents=documents,
        settings=settings,
        embeddings=embeddings,
    )
    progress("semantic", f"Created {len(parents)} parent chunks", len(parents), len(parents))

    all_children: list[Document] = []
    parent_meta: dict[str, dict] = {}

    for i, parent in enumerate(parents):
        parent_id = hashlib.sha256(
            (parent.page_content + str(i)).encode()
        ).hexdigest()[:16]

        parent_meta[parent_id] = {
            "content": parent.page_content,
            "source": parent.metadata.get("source", source_label),
        }

        children = parent_to_children_dynamic(parent, settings, effective_strategy=effective_strategy)
        for child in children:
            for key, value in parent.metadata.items():
                if key not in child.metadata or key == "source":
                    child.metadata[key] = value
            child.metadata["parent_id"] = parent_id
            all_children.append(child)

    upsert_vector_parents(collection_name, parent_meta)

    child_texts = [_sanitize_text_for_api(c.page_content) for c in all_children]
    child_metadatas = []
    for c in all_children:
        m = {k: _sanitize_metadata_value(v) for k, v in c.metadata.items()}
        child_metadatas.append(m)

    child_embeddings = _parallel_embed(
        embeddings,
        child_texts,
        batch_size=getattr(settings, "embed_batch_size", 64),
        max_workers=getattr(settings, "embed_max_workers", 4),
        progress_fn=progress,
    )

    vector_size = len(child_embeddings[0]) if child_embeddings else 0
    if vector_size == 0:
        return {
            "file": source_label,
            "collection_name": collection_name,
            "num_parents": len(parents),
            "num_children": 0,
            "total_chunks_in_db": 0,
        }

    progress("sparse", "Building sparse vocab and vectors...", 0, 1)
    existing_vocab = load_sparse_vocab_map(collection_name)

    cfg = get_settings()
    new_vocab, idf_map, avgdl = _build_vocab(child_texts)
    merged_vocab = dict(existing_vocab)
    next_idx = max(merged_vocab.values(), default=-1) + 1
    added_vocab: dict[str, int] = {}
    for token in new_vocab:
        if token not in merged_vocab:
            merged_vocab[token] = next_idx
            added_vocab[token] = next_idx
            next_idx += 1

    sparse_vectors: list[tuple[list[int], list[float]]] = [
        _text_to_sparse_vector(t, merged_vocab, idf_map, avgdl, cfg.bm25_k1, cfg.bm25_b)
        for t in child_texts
    ]
    insert_sparse_vocab_entries(collection_name, added_vocab)
    upsert_sparse_bm25_stats(collection_name, avgdl, idf_map)
    try:
        from retriever import invalidate_bm25_cache
        invalidate_bm25_cache(collection_name)
    except Exception:
        pass
    progress("sparse", f"Built sparse vocab ({len(merged_vocab)} terms)", 1, 1)

    # Build chunk records for pgvector
    records: list[dict] = []
    for i in range(len(child_texts)):
        indices, values = sparse_vectors[i]
        child = all_children[i]
        parent_id = child.metadata.get("parent_id", "")
        records.append({
            "parent_id": parent_id,
            "chunk_index": int(child.metadata.get("chunk_index", i)),
            "text": child_texts[i],
            "embedding": child_embeddings[i],
            "sparse_indices": indices,
            "sparse_values": values,
            "payload": {**child_metadatas[i], "text": child_texts[i]},
        })

    progress("pgvector", f"Writing {len(records)} chunks to PostgreSQL...", 0, 1)
    _insert_chunks_in_batches(
        collection_name=collection_name,
        records=records,
        batch_size=getattr(settings, "pgvector_upsert_batch_size", 128),
        progress_fn=progress,
    )

    total_in_db = get_collection_count(collection_name)
    progress("done", f"Done: {len(parents)} parents, {len(all_children)} children", 1, 1)

    return {
        "file": source_label,
        "collection_name": collection_name,
        "num_parents": len(parents),
        "num_children": len(all_children),
        "total_chunks_in_db": total_in_db,
    }

def ingest_file(
    file_path: Path,
    *,
    collection_name: str = "rag_chatbot",
    on_progress: Optional[Callable[[str, str, int, int], None]] = None,
) -> dict[str, Any]:
    """Load file, build parent/child chunks, embed and store in PostgreSQL."""
    progress = on_progress or _noop_progress
    collection_name = sanitize_collection_name(collection_name)
    file_path = Path(file_path)
    progress("load", "Loading file...", 0, 1)
    settings = get_settings()
    documents = load_document(file_path, settings=settings)
    return ingest_documents(
        documents,
        source_label=str(file_path),
        collection_name=collection_name,
        on_progress=on_progress,
    )



def ingest_directory(
    dir_path: Path,
    collection_name: str = "rag_chatbot",
    on_progress: Optional[Callable[[str, str, int, int], None]] = None,
) -> list[dict[str, Any]]:
    """Ingest all PDF and DOCX files in a directory."""
    results = []
    for ext in ("*.pdf", "*.docx", "*.doc"):
        for path in Path(dir_path).rglob(ext):
            try:
                r = ingest_file(
                    path,
                    collection_name=collection_name,
                    on_progress=on_progress,
                )
                results.append(r)
            except Exception as e:
                results.append({"file": str(path), "error": str(e)})
    return results
