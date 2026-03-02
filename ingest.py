"""Ingestion: load PDF/DOCX, semantic chunking, parent-child hierarchy, metadata (summary, target_question)."""
import hashlib
import json
import re
from pathlib import Path
from typing import Any, Callable, Optional

import tiktoken
from langchain_core.documents import Document
from langchain_experimental.text_splitter import SemanticChunker
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from qdrant_client import QdrantClient
from qdrant_client.http import models as qmodels

from config import get_settings


# --- Tokenization & text splitting ---

def get_encoding():
    return tiktoken.get_encoding("cl100k_base")


def count_tokens(text: str) -> int:
    return len(get_encoding().encode(text))


def chunk_by_tokens(text: str, max_tokens: int, overlap: int = 50) -> list[str]:
    """Split text into chunks by token count with overlap."""
    enc = get_encoding()
    tokens = enc.encode(text)
    out = []
    start = 0
    while start < len(tokens):
        end = min(start + max_tokens, len(tokens))
        out.append(enc.decode(tokens[start:end]))
        start = end - overlap if end < len(tokens) else len(tokens)
    return out


def _split_paragraphs(text: str) -> list[str]:
    """Split text into paragraphs separated by blank lines."""
    parts = [p.strip() for p in re.split(r"\n\s*\n+", text) if p.strip()]
    if not parts and text.strip():
        parts = [text.strip()]
    return parts


def _split_sentences(text: str) -> list[str]:
    """Simple rule-based sentence splitter (Vietnamese and English)."""
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return []
    parts = re.split(r"(?<=[.!?])\s+", cleaned)
    return [p.strip() for p in parts if p.strip()]


def _split_words(text: str) -> list[str]:
    return text.split()


def split_by_paragraphs(text: str, max_paragraphs: int, overlap: int = 0) -> list[str]:
    paras = _split_paragraphs(text)
    if not paras:
        return []
    max_p = max(1, max_paragraphs)
    chunks: list[str] = []
    start = 0
    while start < len(paras):
        end = min(start + max_p, len(paras))
        chunks.append("\n\n".join(paras[start:end]))
        if end >= len(paras):
            break
        start = max(0, end - max(overlap, 0))
    return chunks


def split_by_sentences(text: str, max_sentences: int, overlap: int = 1) -> list[str]:
    sentences = _split_sentences(text)
    if not sentences:
        return []
    max_s = max(1, max_sentences)
    chunks: list[str] = []
    start = 0
    while start < len(sentences):
        end = min(start + max_s, len(sentences))
        chunks.append(" ".join(sentences[start:end]))
        if end >= len(sentences):
            break
        start = max(0, end - max(overlap, 0))
    return chunks


def split_by_words(text: str, max_words: int, overlap: int = 30) -> list[str]:
    words = _split_words(text)
    if not words:
        return []
    max_w = max(1, max_words)
    chunks: list[str] = []
    start = 0
    while start < len(words):
        end = min(start + max_w, len(words))
        chunks.append(" ".join(words[start:end]))
        if end >= len(words):
            break
        start = max(0, end - max(overlap, 0))
    return chunks


def _estimate_doc_stats(documents: list[Document]) -> dict[str, float]:
    """Compute simple statistics to drive auto chunking decisions."""
    full_text_parts: list[str] = []
    for d in documents:
        if d and getattr(d, "page_content", None):
            full_text_parts.append(str(d.page_content))
    full_text = "\n\n".join(full_text_parts).strip()
    if not full_text:
        return {
            "num_paragraphs": 0,
            "num_sentences": 0,
            "num_words": 0,
            "avg_words_per_paragraph": 0.0,
            "avg_words_per_sentence": 0.0,
        }

    paragraphs = _split_paragraphs(full_text)
    sentences = _split_sentences(full_text)
    words = _split_words(full_text)

    num_paragraphs = len(paragraphs)
    num_sentences = len(sentences)
    num_words = len(words)

    avg_words_per_paragraph = num_words / max(1, num_paragraphs)
    avg_words_per_sentence = num_words / max(1, num_sentences)

    return {
        "num_paragraphs": num_paragraphs,
        "num_sentences": num_sentences,
        "num_words": num_words,
        "avg_words_per_paragraph": avg_words_per_paragraph,
        "avg_words_per_sentence": avg_words_per_sentence,
    }


def auto_select_chunk_strategy(documents: list[Document]) -> str:
    """
    Heuristic to choose a chunking strategy automatically:
    - Prefer paragraph-based when documents are clearly structured into many paragraphs.
    - Prefer sentence-based for long, dense prose.
    - Fall back to semantic token-based when structure is unclear.
    """
    stats = _estimate_doc_stats(documents)
    num_paragraphs = stats["num_paragraphs"]
    num_sentences = stats["num_sentences"]
    avg_words_per_paragraph = stats["avg_words_per_paragraph"]
    avg_words_per_sentence = stats["avg_words_per_sentence"]

    if num_paragraphs >= 10 and 40 <= avg_words_per_paragraph <= 200:
        return "paragraph"
    if num_sentences >= 20 and 10 <= avg_words_per_sentence <= 60:
        return "sentence"
    return "semantic_tokens"


def _load_docx(path: Path) -> list[Document]:
    """Load DOCX via python-docx (no extra system deps)."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        try:
            from langchain_community.document_loaders import UnstructuredWordDocumentLoader
            loader = UnstructuredWordDocumentLoader(str(path))
            return loader.load()
        except Exception:
            raise ImportError("Install python-docx: pip install python-docx")
    doc = DocxDocument(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n\n".join(parts) or ""
    if not text.strip():
        return []
    return [Document(page_content=text, metadata={"source": path.name, "file_path": str(path)})]


def load_document(file_path: Path) -> list[Document]:
    """Load a single PDF or DOCX file into LangChain documents."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        loader = PyPDFLoader(str(path))
        docs = loader.load()
    elif suffix in (".docx", ".doc"):
        docs = _load_docx(path)
    else:
        raise ValueError(f"Unsupported format: {suffix}. Use .pdf or .docx")

    for d in docs:
        d.metadata.setdefault("source", str(path.name))
        d.metadata["file_path"] = str(path)
    return docs


# --- Parent and child chunking ---

def semantic_parent_chunks(
    documents: list[Document],
    parent_target_tokens: int,
    embeddings: Any,
) -> list[Document]:
    """
    Use SemanticChunker to get semantic boundaries, then merge into parent-sized chunks.
    """
    semantic = SemanticChunker(embeddings=embeddings)
    all_splits: list[Document] = []
    for doc in documents:
        splits = semantic.split_documents([doc])
        all_splits.extend(splits)

    parents: list[Document] = []
    current_text: list[str] = []
    current_tokens = 0
    current_metadata: dict = {}

    for d in all_splits:
        t = d.page_content.strip()
        if not t:
            continue
        n = count_tokens(t)
        if current_tokens + n > parent_target_tokens and current_text:
            parent_text = "\n\n".join(current_text)
            parents.append(
                Document(
                    page_content=parent_text,
                    metadata={**current_metadata, **d.metadata},
                )
            )
            current_text = [t]
            current_tokens = n
            current_metadata = d.metadata
        else:
            current_text.append(t)
            current_tokens += n
            current_metadata = d.metadata

    if current_text:
        parent_text = "\n\n".join(current_text)
        parents.append(
            Document(page_content=parent_text, metadata=current_metadata)
        )
    return parents


def paragraph_parent_chunks(
    documents: list[Document],
    parent_paragraphs: int,
) -> list[Document]:
    """Group documents into parent chunks by paragraphs."""
    parents: list[Document] = []
    for doc in documents:
        text = doc.page_content or ""
        paras = _split_paragraphs(text)
        if not paras:
            continue
        current: list[str] = []
        for p in paras:
            current.append(p)
            if len(current) >= max(1, parent_paragraphs):
                parents.append(
                    Document(
                        page_content="\n\n".join(current),
                        metadata=dict(doc.metadata),
                    )
                )
                current = []
        if current:
            parents.append(
                Document(
                    page_content="\n\n".join(current),
                    metadata=dict(doc.metadata),
                )
            )
    return parents


def sentence_parent_chunks(
    documents: list[Document],
    parent_sentences: int,
) -> list[Document]:
    """Group documents into parent chunks by sentence count."""
    parents: list[Document] = []
    for doc in documents:
        text = doc.page_content or ""
        sentences = _split_sentences(text)
        if not sentences:
            continue
        current: list[str] = []
        for s in sentences:
            current.append(s)
            if len(current) >= max(1, parent_sentences):
                parents.append(
                    Document(
                        page_content=" ".join(current),
                        metadata=dict(doc.metadata),
                    )
                )
                current = []
        if current:
            parents.append(
                Document(
                    page_content=" ".join(current),
                    metadata=dict(doc.metadata),
                )
            )
    return parents


def word_parent_chunks(
    documents: list[Document],
    parent_words: int,
) -> list[Document]:
    """Group documents into parent chunks by word count."""
    parents: list[Document] = []
    for doc in documents:
        text = doc.page_content or ""
        words = _split_words(text)
        if not words:
            continue
        max_w = max(1, parent_words)
        start = 0
        while start < len(words):
            end = min(start + max_w, len(words))
            chunk_text = " ".join(words[start:end])
            parents.append(
                Document(
                    page_content=chunk_text,
                    metadata=dict(doc.metadata),
                )
            )
            start = end
    return parents


def build_parent_chunks(
    documents: list[Document],
    settings,
    embeddings: Any,
) -> list[Document]:
    """
    Build parent chunks according to configured strategy:
    - semantic_tokens (default)
    - paragraph
    - sentence
    - word
    - auto (choose best heuristic based on document statistics)
    """
    strategy = getattr(settings, "chunk_strategy", "semantic_tokens") or "semantic_tokens"
    strategy = strategy.lower()

    if strategy == "auto":
        # Use whole-document statistics to pick a concrete strategy
        strategy = auto_select_chunk_strategy(documents)

    if strategy == "paragraph":
        return paragraph_parent_chunks(
            documents,
            parent_paragraphs=getattr(settings, "parent_chunk_paragraphs", 4),
        )
    if strategy == "sentence":
        return sentence_parent_chunks(
            documents,
            parent_sentences=getattr(settings, "parent_chunk_sentences", 8),
        )
    if strategy == "word":
        return word_parent_chunks(
            documents,
            parent_words=getattr(settings, "parent_chunk_words", 400),
        )

    # Fallback to existing semantic + token-based approach
    return semantic_parent_chunks(
        documents,
        parent_target_tokens=getattr(settings, "parent_chunk_tokens", 700),
        embeddings=embeddings,
    )


def parent_to_children_dynamic(parent: Document, settings, *, effective_strategy: Optional[str] = None) -> list[Document]:
    """
    Split a parent chunk into child chunks according to the configured strategy.
    - semantic_tokens: token-based windows
    - paragraph: paragraphs per chunk
    - sentence: sentences per chunk
    - word: words per chunk
    """
    text = parent.page_content or ""
    strategy = effective_strategy or getattr(settings, "chunk_strategy", "semantic_tokens") or "semantic_tokens"
    strategy = strategy.lower()

    if strategy == "auto":
        strategy = auto_select_chunk_strategy([parent])

    if strategy == "paragraph":
        chunks = split_by_paragraphs(
            text,
            max_paragraphs=getattr(settings, "child_chunk_paragraphs", 1),
            overlap=0,
        )
    elif strategy == "sentence":
        chunks = split_by_sentences(
            text,
            max_sentences=getattr(settings, "child_chunk_sentences", 3),
            overlap=1,
        )
    elif strategy == "word":
        chunks = split_by_words(
            text,
            max_words=getattr(settings, "child_chunk_words", 120),
            overlap=30,
        )
    else:
        chunks = chunk_by_tokens(
            text,
            max_tokens=getattr(settings, "child_chunk_tokens", 150),
            overlap=30,
        )

    children: list[Document] = []
    for i, c in enumerate(chunks):
        content = c.strip()
        if not content:
            continue
        children.append(
            Document(
                page_content=content,
                metadata={
                    **parent.metadata,
                    "chunk_index": i,
                    "parent_content": text,
                },
            )
        )
    return children


def add_parent_metadata(parent: Document, llm: ChatOpenAI) -> tuple[str, str]:
    """Generate summary and target_question for a parent chunk via LLM."""
    prompt = """Bạn là trợ lý tóm tắt. Cho đoạn văn sau (có thể bằng tiếng Việt hoặc tiếng Anh), hãy:
1. Tóm tắt ngắn gọn nội dung chính (2-3 câu) bằng cùng ngôn ngữ với tài liệu.
2. Viết một câu hỏi mẫu mà đoạn văn này có thể trả lời (target question).

Đoạn văn:
---
{text}
---

Trả lời theo đúng format sau (giữ nguyên nhãn):
SUMMARY: <tóm tắt>
TARGET_QUESTION: <câu hỏi mẫu>"""
    msg = prompt.format(text=parent.page_content[:4000])
    try:
        resp = llm.invoke(msg)
        content = resp.content if hasattr(resp, "content") else str(resp)
        summary = ""
        target_question = ""
        for line in content.strip().split("\n"):
            if line.upper().startswith("SUMMARY:"):
                summary = line.split(":", 1)[1].strip()
            elif line.upper().startswith("TARGET_QUESTION:"):
                target_question = line.split(":", 1)[1].strip()
        return summary or content[:200], target_question or "Nội dung đoạn văn là gì?"
    except Exception:
        return parent.page_content[:200], "Nội dung đoạn văn là gì?"


# --- Collection name & Qdrant client ---

def sanitize_collection_name(name: str) -> str:
    """Normalize collection name: 3-63 chars, alphanumeric/underscore/hyphen."""
    if not name or not str(name).strip():
        return "rag_chatbot"
    s = str(name).strip()
    s = re.sub(r"[^a-zA-Z0-9_-]", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    if not s:
        return "rag_chatbot"
    if len(s) < 3:
        s = s + "xx"[: 3 - len(s)]
    if not s[0].isalnum():
        s = "c" + s
    if not s[-1].isalnum():
        s = s + "1"
    return s[:63]


def get_qdrant_client(persist_dir: Path) -> QdrantClient:
    """Create or connect to local Qdrant under persist_dir."""
    persist_dir.mkdir(parents=True, exist_ok=True)
    path = persist_dir / "storage"
    path.mkdir(parents=True, exist_ok=True)
    return QdrantClient(path=str(path))


def _noop_progress(step: str, msg: str, current: int = 0, total: int = 0) -> None:
    pass


# --- Full file ingestion ---

def ingest_file(
    file_path: Path,
    *,
    collection_name: str = "rag_chatbot",
    skip_metadata_llm: bool = False,
    on_progress: Optional[Callable[[str, str, int, int], None]] = None,
) -> dict[str, Any]:
    """
    Load file, build parent/child chunks, add metadata, store in Qdrant.
    Returns dict with num_parents, num_children, etc. on_progress(step, msg, current, total) is optional.
    """
    progress = on_progress or _noop_progress
    collection_name = sanitize_collection_name(collection_name)
    settings = get_settings()
    if not settings.openai_api_key:
        raise ValueError("OPENAI_API_KEY is required. Set it in .env or environment.")

    file_path = Path(file_path)
    progress("load", "Loading file...", 0, 1)
    documents = load_document(file_path)
    if not documents:
        return {"error": "No content extracted", "file": str(file_path)}
    progress("load", f"Loaded {len(documents)} page(s)", 1, 1)

    if settings.embedding_model.startswith("text-embedding"):
        embeddings = OpenAIEmbeddings(
            model=settings.embedding_model,
            api_key=settings.openai_api_key,
        )
    else:
        embeddings = HuggingFaceEmbeddings(
            model_name=settings.embedding_model,
        )
    llm = ChatOpenAI(
        model=settings.llm_model,
        api_key=settings.openai_api_key,
        temperature=0.2,
    )
    configured_strategy = getattr(settings, "chunk_strategy", "semantic_tokens") or "semantic_tokens"
    configured_strategy = configured_strategy.lower()
    if configured_strategy == "auto":
        effective_strategy = auto_select_chunk_strategy(documents)
        progress("semantic", f"Auto chunking (strategy: {effective_strategy})...", 0, 0)
    else:
        effective_strategy = configured_strategy
        progress("semantic", f"Chunking ({effective_strategy})...", 0, 0)

    parents = build_parent_chunks(
        documents=documents,
        settings=settings,
        embeddings=embeddings,
    )
    progress("semantic", f"Created {len(parents)} parent chunks", len(parents), len(parents))

    all_children: list[Document] = []
    parent_meta: dict[str, dict] = {}

    for i, parent in enumerate(parents):
        if not skip_metadata_llm:
            progress("metadata", f"Summary for parent {i + 1}/{len(parents)}...", i + 1, len(parents))
        parent_id = hashlib.sha256(
            (parent.page_content + str(i)).encode()
        ).hexdigest()[:16]
        summary = ""
        target_question = ""
        if not skip_metadata_llm:
            summary, target_question = add_parent_metadata(parent, llm)
        else:
            summary = parent.page_content[:200]
            target_question = "Nội dung đoạn văn là gì?"

        parent_meta[parent_id] = {
            "content": parent.page_content,
            "summary": summary,
            "target_question": target_question,
            "source": parent.metadata.get("source", file_path.name),
        }

        children = parent_to_children_dynamic(parent, settings, effective_strategy=effective_strategy)
        for j, child in enumerate(children):
            child.metadata["parent_id"] = parent_id
            child.metadata["summary"] = summary
            child.metadata["target_question"] = target_question
            child.metadata["parent_content"] = parent.page_content
            child.metadata["source"] = parent.metadata.get("source", file_path.name)
            all_children.append(child)

    settings.persist_dir.mkdir(parents=True, exist_ok=True)
    client = get_qdrant_client(settings.persist_dir)

    parents_path = settings.persist_dir / f"{collection_name}_parents.json"
    with open(parents_path, "w", encoding="utf-8") as f:
        json.dump(parent_meta, f, ensure_ascii=False, indent=2)

    child_texts = [c.page_content for c in all_children]
    child_metadatas = []
    for c in all_children:
        m = {k: (v if isinstance(v, (str, int, float, bool)) else str(v)) for k, v in c.metadata.items()}
        if len(m.get("parent_content", "")) > 30000:
            m["parent_content"] = m["parent_content"][:30000] + "..."
        child_metadatas.append(m)

    progress("embed", f"Embedding {len(child_texts)} chunks...", 0, 1)
    child_embeddings = embeddings.embed_documents(child_texts)
    progress("embed", f"Embedded {len(child_embeddings)} chunks", 1, 1)

    vector_size = len(child_embeddings[0]) if child_embeddings else 0
    if vector_size == 0:
        return {
            "file": str(file_path),
            "collection_name": collection_name,
            "num_parents": len(parents),
            "num_children": 0,
            "total_chunks_in_db": 0,
        }

    try:
        client.get_collection(collection_name)
    except Exception:
        client.create_collection(
            collection_name=collection_name,
            vectors_config=qmodels.VectorParams(
                size=vector_size,
                distance=qmodels.Distance.COSINE,
            ),
        )

    from qdrant_client.http import models as qmodels_local

    ids = list(range(len(child_texts)))
    points = [
        qmodels_local.PointStruct(
            id=ids[i],
            vector=child_embeddings[i],
            payload={**child_metadatas[i], "text": child_texts[i]},
        )
        for i in range(len(child_texts))
    ]

    progress("qdrant", f"Writing {len(points)} chunks to Qdrant...", 0, 1)
    client.upsert(collection_name=collection_name, points=points)

    existing = client.count(collection_name=collection_name, exact=True).count
    progress("done", f"Done: {len(parents)} parents, {len(all_children)} children", 1, 1)

    return {
        "file": str(file_path),
        "collection_name": collection_name,
        "num_parents": len(parents),
        "num_children": len(all_children),
        "total_chunks_in_db": existing + len(all_children),
    }


def ingest_directory(
    dir_path: Path,
    collection_name: str = "rag_chatbot",
    skip_metadata_llm: bool = False,
    on_progress: Optional[Callable[[str, str, int, int], None]] = None,
) -> list[dict[str, Any]]:
    """Ingest all PDF and DOCX files in a directory."""
    results = []
    for ext in ("*.pdf", "*.docx", "*.doc"):
        for path in Path(dir_path).rglob(ext):
            try:
                r = ingest_file(
                    path,
                    collection_name=collection_name,
                    skip_metadata_llm=skip_metadata_llm,
                    on_progress=on_progress,
                )
                results.append(r)
            except Exception as e:
                results.append({"file": str(path), "error": str(e)})
    return results
